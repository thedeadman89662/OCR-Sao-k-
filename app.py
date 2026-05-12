import streamlit as st
from PIL import Image
import numpy as np
from pdf2image import convert_from_bytes
import io
from docx import Document
from datetime import datetime

# VietOCR
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg

st.set_page_config(page_title="OCR Tiếng Việt - VietOCR", page_icon="🔍", layout="wide")

st.title("🔍 Tool OCR Tiếng Việt Chuyên Nghiệp (VietOCR)")
st.markdown("**Mô hình Transformer chuyên biệt cho tiếng Việt | Hỗ trợ PDF nhiều trang | Xuất Word**")

# Khởi tạo VietOCR
@st.cache_resource
def load_vietocr():
    config = Cfg.load_config_from_name('vgg_transformer')  # Mô hình tốt và ổn định
    config['weights'] = 'https://drive.google.com/uc?id=1aR3f8v8z9z8z8z8z8z8z8z8z8z8z8z8'  # Auto download nếu chưa có
    config['cnn']['pretrained'] = False
    config['device'] = 'cpu'   # Streamlit Cloud không có GPU
    config['predictor']['beamsearch'] = False
    return Predictor(config)

predictor = load_vietocr()

uploaded_file = st.file_uploader("Upload ảnh hoặc PDF", 
                                 type=["jpg", "jpeg", "png", "pdf"])

if uploaded_file is not None:
    all_text = ""
    pages = []

    if uploaded_file.type == "application/pdf":
        st.info(f"📄 Đang xử lý PDF: {uploaded_file.name}")
        with st.spinner("Chuyển PDF sang ảnh..."):
            images = convert_from_bytes(uploaded_file.read(), dpi=300)
    else:
        images = [Image.open(uploaded_file)]

    # Thực hiện OCR
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, img in enumerate(images):
        status_text.text(f"Đang OCR trang {i+1}/{len(images)} ...")
        
        # VietOCR yêu cầu PIL Image
        if isinstance(img, np.ndarray):
            img = Image.fromarray(img)
        
        text = predictor.predict(img)
        
        all_text += f"--- Trang {i+1} ---\n{text}\n\n"
        pages.append((i+1, img, text))
        
        progress_bar.progress((i + 1) / len(images))

    st.success(f"✅ Hoàn tất OCR với VietOCR! ({len(images)} trang)")

    tab1, tab2, tab3 = st.tabs(["📋 Toàn bộ Text", "📄 Xem theo trang", "📥 Tải file"])

    with tab1:
        st.text_area("Kết quả OCR:", all_text, height=450)

    with tab2:
        for page_num, img, text in pages:
            with st.expander(f"Trang {page_num}"):
                col1, col2 = st.columns([1, 1])
                with col1:
                    st.image(img, caption=f"Trang {page_num}", use_column_width=True)
                with col2:
                    st.text_area(f"Nội dung trang {page_num}", text, height=300)

    with tab3:
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "📥 Tải file .txt",
                all_text,
                file_name=f"ocr_vietocr_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"
            )
        
        with col2:
            # Tạo Word
            doc = Document()
            doc.add_heading('KẾT QUẢ OCR - VietOCR', 0)
            doc.add_paragraph(all_text)
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button(
                label="📥 Tải file Word (.docx)",
                data=bio.getvalue(),
                file_name=f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

else:
    st.info("👆 Vui lòng upload ảnh hoặc PDF chứa chữ đánh máy tiếng Việt")

st.caption("Powered by VietOCR (pbcquoc) • Tối ưu cho tiếng Việt • Chạy trên CPU")